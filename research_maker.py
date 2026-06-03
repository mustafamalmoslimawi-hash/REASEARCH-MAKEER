import streamlit as st
import requests
import docx
from io import BytesIO
import xml.etree.ElementTree as ET
import google.generativeai as genai

# إعدادات واجهة المستخدم الرسومية لـ RESEARCH-MAKER ULTRA
st.set_page_config(page_title="RESEARCH-MAKER ULTRA", layout="wide")

st.markdown("<h1 style='text-align: center; color: #008080;'>🔬 RESEARCH-MAKER ULTRA</h1>", unsafe_allow_html=True)
st.markdown("<h4 style='text-align: center; color: #555;'>المحرك الأكاديمي الشامل: جلب متعدد المصادر + محاكاة قالب الـ Word المرفوع</h4>", unsafe_allow_html=True)
st.write("---")

# جلب المفاتيح بأمان كامل من خزنة Streamlit (Secrets) لضمان عدم حظرها
SERPAPI_KEY = st.secrets.get("SERPAPI_KEY", "").strip()
GEMINI_KEY = st.secrets.get("GEMINI_KEY", "").strip()

# تهيئة مكتبة الجيل الجديد لـ Gemini بالمفتاح السري
if GEMINI_KEY:
    genai.configure(api_key=GEMINI_KEY)

# ==================== قسم قراءة قالب الـ WORD ====================
def extract_structure_from_docx(file_buffer):
    """قراءة مستند الـ Word المرفوع واستخراج الهيكل والترتيب منه تلقائياً"""
    try:
        doc = docx.Document(file_buffer)
        structure_lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if para.style.name.startswith('Heading') or text.isupper() or any(k in text.lower() for k in ['chapter', 'section', 'المبحث', 'الفصل']):
                    structure_lines.append(f"[Heading] {text}")
                else:
                    structure_lines.append(text)
        return "\n".join(structure_lines)
    except Exception as e:
        st.error(f"حدث خطأ أثناء تحليل ملف القالب: {e}")
        return ""

# ==================== قسم الـ APIs الموسعة لتجهيز البحث ====================
def fetch_semantic_scholar(query):
    """[API 1] سحب الأبحاث الأكاديمية مجاناً من Semantic Scholar وبدون حدود"""
    url = f"https://api.semanticscholar.org/graph/v1/paper/search?query={query}&limit=8&fields=title,abstract,url"
    try:
        response = requests.get(url, timeout=10)
        data = response.json().get("data", [])
        papers = []
        for item in data:
            if item.get("abstract") or item.get("title"):
                papers.append({
                    "title": item.get("title", "No Title"),
                    "snippet": item.get("abstract", "No abstract available."),
                    "link": item.get("url", "#"),
                    "source": "Semantic Scholar"
                })
        return papers
    except:
        return []

def fetch_arxiv(query):
    """[API 2] سحب الأوراق الأكاديمية المفتوحة من مستودع arXiv العالمي مجاناً"""
    url = f"http://export.arxiv.org/api/query?search_query=all:{query}&max_results=8"
    try:
        response = requests.get(url, timeout=10)
        root = ET.fromstring(response.text)
        papers = []
        for entry in root.findall('{http://www.w3.org/2005/Atom}entry'):
            title = entry.find('{http://www.w3.org/2005/Atom}title').text
            summary = entry.find('{http://www.w3.org/2005/Atom}summary')
            snippet = summary.text if summary is not None else "No summary available."
            link = entry.find('{http://www.w3.org/2005/Atom}id').text
            papers.append({
                "title": title.strip().replace("\n", ""),
                "snippet": snippet.strip().replace("\n", " "),
                "link": link,
                "source": "arXiv Repository"
            })
        return papers
    except:
        return []

def fetch_google_scholar(query):
    """[API 3] جلب الأبحاث العلمية عبر SerpApi"""
    if not SERPAPI_KEY: return []
    url = "https://serpapi.com/search"
    params = {"engine": "google_scholar", "q": query, "hl": "en", "api_key": SERPAPI_KEY}
    try:
        response = requests.get(url, params=params, timeout=10)
        results = response.json().get("organic_results", [])
        return [{
            "title": item.get("title", "No Title"),
            "snippet": item.get("snippet", "No abstract available."),
            "link": item.get("link", "#"),
            "source": "Google Scholar"
        } for item in results[:8]]
    except:
        return []

# ==================== عقل التوليد الذكي وصناعة المحتوى ====================
def generate_advanced_templated_research(title, combined_papers, template_text, lang, style):
    """توجيه ذكاء Gemini لبناء محتوى البحث بالاعتماد الكلي على هيكل قالب الورد والمصادر الموسعة"""
    if not GEMINI_KEY:
        return "ERROR_KEY: لم يتم إضافة مفتاح Gemini API بشكل صحيح في خزنة الـ Secrets لـ Streamlit."
        
    sources_block = ""
    for idx, p in enumerate(combined_papers):
        sources_block += f"\n[Source {idx+1} from {p['source']}]\nTitle: {p['title']}\nAbstract: {p['snippet']}\nLink: {p['link']}\n"
        
    prompt = f"""
    You are an elite academic professor. Write an extensive, deep, and fully cited scientific research paper about the new topic: '{title}'.
    
    CRITICAL STRUCTURE MANDATE: You must completely mimic, follow, and fill the exact section flow, architectural pattern, and outline sequence from the "UPLOADED TEMPLATE STRUCT" provided below. Maintain its structure completely but generate entirely new academic content for '{title}'.
    
    Target Language: {lang}
    Citation Style: {style}
    
    [UPLOADED TEMPLATE STRUCT - FOLLOW THIS EXACTLY]:
    {template_text}
    
    [ACADEMIC DATABASE TO BUILD FROM]:
    {sources_block}
    
    Execution Plan:
    1. Fill each title/chapter from the template structure with rich, advanced, and elongated technical analysis.
    2. Inject extensive in-text citations linking to the sources provided.
    3. Make sure to generate all sections comprehensively to give maximum depth.
    4. Compile a perfect references grid at the end based on {style}.
    """
    
    try:
        model = genai.GenerativeModel('gemini-1.5-pro')
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                max_output_tokens=8192,
                temperature=0.3
            )
        )
        return response.text
    except Exception as e:
        return f"ERROR_API: {e}"

def create_formatted_docx(text, title):
    """حفظ وتصدير البحث العلمي الجديد في مستند Word منظم"""
    doc = docx.Document()
    doc.add_heading(title, level=0)
    lines = text.split("\n")
    for line in lines:
        clean_line = line.replace("**", "").replace("###", "").replace("##", "").replace("#", "").strip()
        if not clean_line: continue
        if any(keyword in line for keyword in ["CHAPTER", "Introduction", "Review", "Discussion", "Conclusion", "References", "المبحث", "الفصل"]):
            doc.add_heading(clean_line, level=1)
        else:
            doc.add_paragraph(clean_line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==================== بناء واجهة المستخدم الرسومية المستقرة ====================
research_title = st.text_input("📝 أولاً: اكتب عنوان البحث العلمي الجديد المُراد صناعته:")
uploaded_file = st.file_uploader("📐 ثانياً: ارفع ملف الـ Word القياسي (ليقوم النظام بنسخ هيكليته وفصوله التنسيقية):", type=["docx"])

col1, col2 = st.columns(2)
with col1:
    language = st.selectbox("🌐 لغة الكتابة الأكاديمية المطلوبة:", ["English", "العربية"])
with col2:
    citation_style = st.selectbox("📚 نظام توثيق الهوامش والملحقات:", ["APA", "IEEE", "Harvard"])

if st.button("🚀 تشغيل النظام الفائق وإنشاء البحث"):
    if research_title and uploaded_file is not None:
        with st.spinner("📊 جاري فتح مستند الـ Word وقراءة القالب البنيوي له..."):
            extracted_template = extract_structure_from_docx(uploaded_file)
            
        if extracted_template:
            with st.spinner("🌐 جاري جلب الأوراق والمراجع من المستودعات الأكاديمية العالمية..."):
                res_google = fetch_google_scholar(research_title)
                res_semantic = fetch_semantic_scholar(research_title)
                res_arxiv = fetch_arxiv(research_title)
                
                all_combined_papers = res_google + res_semantic + res_arxiv
                
            if not all_combined_papers:
                all_combined_papers = [{
                    "title": f"General Overview on {research_title}",
                    "snippet": f"Academic background and foundational concepts regarding {research_title}.",
                    "link": "https://scholar.google.com",
                    "source": "Local Academic Framework"
                }]
                
            st.success(f"🔥 تم تجميع وتأمين {len(all_combined_papers)} مرجعاً علمياً متقاطعة لتوثيق بحثك!")
            
            with st.expander("🔗 استكشاف قائمة المراجع الشاملة المجهّزة للبحث"):
                for idx, paper in enumerate(all_combined_papers):
                    st.markdown(f"**[{idx+1}] {paper['title']}** — <span style='color:green;'>{paper['source']}</span>", unsafe_allow_html=True)
                    st.write(paper['snippet'])
                    st.markdown(f"[رابط المصدر الدائم]({paper['link']})")
                    st.write("---")
                    
            with st.spinner("🧠 يقوم نظام Gemini 1.5 Pro الآن بمطابقة هيكلية ملف الورد المرفوع وصياغة الفصول كاملة..."):
                generated_research = generate_advanced_templated_research(research_title, all_combined_papers, extracted_template, language, citation_style)
                
            st.subheader("📄 معاينة البحث الهيكلي الجديد المولد:")
            st.text_area("المستند الأكاديمي الكامل", generated_research, height=450)
            
            # تم تعديل الشرط هنا ليعمل بنجاح دون أخطاء نصية تفوت تفعيل الأزرار
            if "ERROR_" not in generated_research:
                st.balloons()
                final_docx = create_formatted_docx(generated_research, research_title)
                st.download_button(
                    label="📥 تحميل مستند البحث العلمي الكامل المنسق تلقائياً (.docx)",
                    data=final_docx,
                    file_name=f"{research_title.replace(' ', '_')}_Final_Research.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error(f"تعذر إنتاج التقرير بسبب مشكلة في الخادم الرئيسي: {generated_research}")
    else:
        st.error("الرجاء إدخال عنوان البحث الجديد ورفع ملف قالب الـ Word أولاً لتشغيل النظام الجديد.")
